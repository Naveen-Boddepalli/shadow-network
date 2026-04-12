# ai-engine/extractor.py
# Extracts plain text from various file formats.
# Input: raw bytes + MIME type string
# Output: plain text string

import io
import fitz          # PyMuPDF — handles PDFs
import docx          # python-docx — handles .docx files


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    """
    Dispatch to the right extractor based on MIME type.
    Returns plain text, or raises ValueError for unsupported types.
    """
    if mime_type == "application/pdf":
        return _extract_pdf(file_bytes)

    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_docx(file_bytes)

    elif mime_type == "text/plain":
        return _extract_txt(file_bytes)

    else:
        raise ValueError(f"Unsupported file type for text extraction: {mime_type}")


def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from a PDF using PyMuPDF (fitz).
    Works on text-based PDFs. Scanned image PDFs will return empty text
    (OCR support can be added in Phase 5 using pytesseract).
    """
    text_parts = []

    # Open PDF from bytes (no temp file needed)
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")  # "text" mode = plain text, no HTML
            if page_text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

    full_text = "\n\n".join(text_parts)

    if not full_text.strip():
        raise ValueError(
            "No text found in PDF. It may be a scanned image — OCR not yet supported."
        )

    return full_text


def _extract_docx(file_bytes: bytes) -> str:
    """
    Extract all paragraphs from a .docx file using python-docx.
    Includes table cell text as well.
    """
    file_stream = io.BytesIO(file_bytes)
    doc = docx.Document(file_stream)

    text_parts = []

    # Regular paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Table cells — academic docs often have data in tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    full_text = "\n".join(text_parts)

    if not full_text.strip():
        raise ValueError("No text found in DOCX file.")

    return full_text


def _extract_txt(file_bytes: bytes) -> str:
    """
    Decode plain text, trying UTF-8 first, falling back to latin-1.
    """
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def chunk_text(text: str, max_chars: int = 3000) -> list[str]:
    """
    Split long text into overlapping chunks for model processing.
    Models have a token limit — we chunk large docs and process each chunk.

    Strategy: split on paragraph boundaries, not mid-sentence.
    """
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    current_chunk = []
    current_len = 0

    for para in paragraphs:
        if current_len + len(para) > max_chars and current_chunk:
            chunks.append("\n".join(current_chunk))
            # Overlap: keep last 2 paragraphs in next chunk for context
            current_chunk = current_chunk[-2:]
            current_len = sum(len(p) for p in current_chunk)

        current_chunk.append(para)
        current_len += len(para)

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks
